import sys
import os
from datetime import datetime
from PyQt6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, 
                            QHBoxLayout, QListWidget, QListWidgetItem, QTextEdit, 
                            QPushButton, QLabel, QScrollArea, QSplitter, QComboBox,
                            QMenuBar, QMenu, QFileDialog, QMessageBox)
from PyQt6.QtCore import Qt, QPoint, QRect, pyqtSignal, QTimer
from PyQt6.QtGui import QPixmap, QScreen, QPainter, QAction
from pynput import mouse
from docx import Document
from docx.shared import Inches
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import tempfile

class ScreenshotManager(QMainWindow):
    screenshot_taken = pyqtSignal(QPixmap)
    
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Screenshot Manager")
        self.screenshots = []
        self.is_recording = False
        self.mouse_listener = None
        self.capture_width = 800
        self.capture_height = 600
        self.current_screen = None
        self.preview_timer = QTimer()
        self.preview_timer.timeout.connect(self.update_live_preview)
        self.preview_timer.start(100)  # Update every 100ms
        self.current_screenshot_index = None  # Track current screenshot index
        self.initUI()
        self.createMenuBar()
        
        # Disable record button by default since no screen is selected
        self.record_button.setEnabled(False)
        
        self.screenshot_taken.connect(self.add_screenshot)
        
    def initUI(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        layout = QHBoxLayout(main_widget)
        
        # Create splitter
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # Left panel
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        
        # Control buttons and screen selection
        control_layout = QHBoxLayout()
        self.record_button = QPushButton("Start Recording")
        self.record_button.clicked.connect(self.toggle_recording)
        
        # Screen selection combo box
        self.screen_combo = QComboBox()
        self.screen_combo.addItem("None")
        for i, screen in enumerate(QApplication.screens()):
            self.screen_combo.addItem(f"Screen {i + 1}")
        self.screen_combo.currentIndexChanged.connect(self.screen_selected)
        
        control_layout.addWidget(self.record_button)
        control_layout.addWidget(QLabel("Select Screen:"))
        control_layout.addWidget(self.screen_combo)
        left_layout.addLayout(control_layout)
        
        # Live preview
        self.live_preview_label = QLabel("Live Preview")
        self.live_preview_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.live_preview_label.setMinimumSize(400, 300)
        left_layout.addWidget(self.live_preview_label)
        
        # Screenshot list
        self.list_widget = QListWidget()
        self.list_widget.itemClicked.connect(self.show_screenshots)
        self.list_widget.itemSelectionChanged.connect(self.on_selection_changed)
        self.list_widget.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.list_widget.customContextMenuRequested.connect(self.show_context_menu)
        left_layout.addWidget(QLabel("Screenshots:"))
        left_layout.addWidget(self.list_widget)
        
        # Right panel
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        # Vertical scroll area for stacked screenshots
        self.preview_scroll = QScrollArea()
        self.preview_container = QWidget()
        self.preview_layout = QVBoxLayout(self.preview_container)
        self.preview_scroll.setWidget(self.preview_container)
        self.preview_scroll.setWidgetResizable(True)
        right_layout.addWidget(QLabel("Preview:"))
        right_layout.addWidget(self.preview_scroll)
        
        # Notes area
        self.notes_edit = QTextEdit()
        right_layout.addWidget(QLabel("Notes:"))
        right_layout.addWidget(self.notes_edit)
        
        # Add panels to splitter
        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        
        # Add splitter to main layout
        layout.addWidget(splitter)
        
        self.setGeometry(100, 100, 1600, 1000)  # Larger window
        
    def screen_selected(self, index):
        if index == 0:  # None selected
            self.current_screen = None
            self.record_button.setEnabled(False)
            self.live_preview_label.clear()
            self.live_preview_label.setText("Live Preview")
        else:
            self.current_screen = QApplication.screens()[index - 1]
            self.record_button.setEnabled(True)
    
    def update_live_preview(self):
        if self.current_screen and not self.is_recording:
            screen_geo = self.current_screen.geometry()
            # Use WinId 0 for the entire desktop and adjust coordinates relative to the screen
            preview = self.current_screen.grabWindow(
                0,
                0,  # Relative to the screen
                0,
                screen_geo.width(),
                screen_geo.height()
            )
            # Scale the preview to fit the label while maintaining aspect ratio
            scaled_preview = preview.scaled(
                self.live_preview_label.size(),
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
            self.live_preview_label.setPixmap(scaled_preview)
    
    def toggle_recording(self):
        if not self.current_screen:
            return
        if not self.is_recording:
            self.is_recording = True
            self.record_button.setText("Stop Recording")
            self.record_button.setStyleSheet("background-color: #ff6b6b;")
            self.mouse_listener = mouse.Listener(on_click=self._handle_click)
            self.mouse_listener.start()
        else:
            self.is_recording = False
            self.record_button.setText("Start Recording")
            self.record_button.setStyleSheet("")
            if self.mouse_listener:
                self.mouse_listener.stop()
                self.mouse_listener = None
    
    def get_screen_at_position(self, x, y):
        """Get the screen containing the given coordinates."""
        point = QPoint(x, y)
        screens = QApplication.screens()
        
        for screen in screens:
            geometry = screen.geometry()
            if geometry.contains(point):
                return screen, geometry
                
        return QApplication.primaryScreen(), QApplication.primaryScreen().geometry()
    
    def _handle_click(self, x, y, button, pressed):
        if button == mouse.Button.left and pressed and self.is_recording:
            screen_geo = self.current_screen.geometry()
            point = QPoint(x, y)
            
            if screen_geo.contains(point):
                # Calculate capture region centered on click
                half_width = self.capture_width // 2
                half_height = self.capture_height // 2
                
                # Adjust coordinates relative to screen
                screen_x = x - screen_geo.x()
                screen_y = y - screen_geo.y()
                
                # Calculate region
                region = QRect(
                    screen_x - half_width,
                    screen_y - half_height,
                    self.capture_width,
                    self.capture_height
                )
                
                # Ensure region stays within screen bounds
                if region.left() < 0:
                    region.moveLeft(0)
                if region.top() < 0:
                    region.moveTop(0)
                if region.right() > screen_geo.width():
                    region.moveRight(screen_geo.width())
                if region.bottom() > screen_geo.height():
                    region.moveBottom(screen_geo.height())
                
                # Capture screenshot using coordinates relative to the screen
                pixmap = self.current_screen.grabWindow(
                    0,
                    region.x(),  # Use coordinates relative to the screen
                    region.y(),
                    region.width(),
                    region.height()
                )
                
                self.screenshot_taken.emit(pixmap)

    
    def add_screenshot(self, pixmap):
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        self.screenshots.append({
            'pixmap': pixmap,
            'timestamp': timestamp,
            'notes': ''
        })
        
        self.list_widget.addItem(timestamp)
        self.list_widget.setCurrentRow(self.list_widget.count() - 1)
        self.show_screenshots(self.list_widget.currentItem())
        
    def on_selection_changed(self):
        # Save current notes before changing selection
        if self.current_screenshot_index is not None and self.current_screenshot_index < len(self.screenshots):
            self.save_current_notes()
        
        # Update for new selection
        selected_items = self.list_widget.selectedItems()
        if selected_items:
            self.show_screenshots(selected_items[0])
        else:
            self.notes_edit.clear()
            self.current_screenshot_index = None
    
    def show_screenshots(self, current_item=None):
        # Clear existing previews
        while self.preview_layout.count():
            item = self.preview_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()
        
        # Show only the selected screenshot
        if current_item:
            current_index = self.list_widget.row(current_item)
            self.current_screenshot_index = current_index  # Update current index
            screenshot = self.screenshots[current_index]
            
            # Create and configure preview label
            preview_label = QLabel()
            preview_label.setPixmap(screenshot['pixmap'])
            preview_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            
            # Scale the preview to fit the scroll area while maintaining aspect ratio
            scaled_pixmap = screenshot['pixmap'].scaled(
                self.preview_scroll.size(),
                Qt.AspectRatioMode.KeepAspectRatio,
                Qt.TransformationMode.SmoothTransformation
            )
            preview_label.setPixmap(scaled_pixmap)
            self.preview_layout.addWidget(preview_label)
            
            # Add timestamp label
            timestamp_label = QLabel(f"Screenshot {screenshot['timestamp']}")
            timestamp_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.preview_layout.addWidget(timestamp_label)
            
            # Update notes
            self.notes_edit.setText(screenshot['notes'])
        else:
            self.current_screenshot_index = None
            self.notes_edit.clear()
        
        # Add stretch at the end
        self.preview_layout.addStretch()

    
    def save_current_notes(self):
        """Save notes for the current screenshot."""
        if (self.current_screenshot_index is not None and 
            0 <= self.current_screenshot_index < len(self.screenshots)):
            current_notes = self.notes_edit.toPlainText()
            self.screenshots[self.current_screenshot_index]['notes'] = current_notes
            
            # Update list item to show note indicator
            list_item = self.list_widget.item(self.current_screenshot_index)
            if list_item:
                timestamp = self.screenshots[self.current_screenshot_index]['timestamp']
                list_item.setText(f"{timestamp} {'📝' if current_notes else ''}")

    def update_notes(self, index):
        """Update notes in real-time."""
        if 0 <= index < len(self.screenshots):
            current_notes = self.notes_edit.toPlainText()
            self.screenshots[index]['notes'] = current_notes
            
            # Update list item to show note indicator
            list_item = self.list_widget.item(index)
            if list_item:
                timestamp = self.screenshots[index]['timestamp']
                list_item.setText(f"{timestamp} {'📝' if current_notes else ''}")
    
    def closeEvent(self, event):
        if self.mouse_listener:
            self.mouse_listener.stop()
        event.accept()

    def createMenuBar(self):
        menubar = self.menuBar()
        file_menu = menubar.addMenu('File')
        
        # Export as Word
        export_word_action = QAction('Export as Word', self)
        export_word_action.triggered.connect(self.export_as_word)
        file_menu.addAction(export_word_action)
        
        # Export as PDF
        export_pdf_action = QAction('Export as PDF', self)
        export_pdf_action.triggered.connect(self.export_as_pdf)
        file_menu.addAction(export_pdf_action)

    def export_as_word(self):
        if not self.screenshots:
            QMessageBox.warning(self, "Warning", "No screenshots to export!")
            return
            
        file_name, _ = QFileDialog.getSaveFileName(self, "Save Word Document", "", 
                                                 "Word Documents (*.docx)")
        if file_name:
            doc = Document()
            doc.add_heading('SOG', 0)
            
            for i, screenshot in enumerate(self.screenshots):
                # Save screenshot to temporary file
                temp_image = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                screenshot['pixmap'].save(temp_image.name, 'PNG')
                
                # Add image number instead of timestamp
                doc.add_heading(f"Image {i + 1}", level=1)
                
                # Add image
                doc.add_picture(temp_image.name, width=Inches(6))
                
                # Add notes with "Comment:" prefix
                if screenshot['notes']:
                    doc.add_paragraph(f"Comment: {screenshot['notes']}")
                
                doc.add_paragraph('\n')  # Add spacing
                
                # Clean up temp file
                temp_image.close()
                os.unlink(temp_image.name)
            
            try:
                doc.save(file_name)
                QMessageBox.information(self, "Success", "Document exported successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save document: {str(e)}")

    def export_as_pdf(self):
        if not self.screenshots:
            QMessageBox.warning(self, "Warning", "No screenshots to export!")
            return
            
        file_name, _ = QFileDialog.getSaveFileName(self, "Save PDF Document", "", 
                                                 "PDF Documents (*.pdf)")
        if file_name:
            try:
                c = canvas.Canvas(file_name, pagesize=letter)
                width, height = letter
                
                # Add title on first page
                c.setFont("Helvetica-Bold", 24)
                c.drawString(50, height - 50, "SOG")
                
                for i, screenshot in enumerate(self.screenshots):
                    if i > 0:
                        c.showPage()  # New page for each screenshot
                    
                    # Save screenshot to temporary file
                    temp_image = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
                    screenshot['pixmap'].save(temp_image.name, 'PNG')
                    
                    # Add image number
                    c.setFont("Helvetica-Bold", 14)
                    c.drawString(50, height - 50, f"Image {i + 1}")
                    
                    # Add image
                    c.drawImage(temp_image.name, 50, height - 400, width=500, preserveAspectRatio=True)
                    
                    # Add notes with "Comment:" prefix
                    if screenshot['notes']:
                        c.setFont("Helvetica", 12)
                        # Split notes into lines to avoid text overflow
                        y_position = height - 450
                        c.drawString(50, y_position, "Comment:")
                        y_position -= 20
                        for line in screenshot['notes'].split('\n'):
                            if y_position > 50:  # Ensure we don't write below page
                                c.drawString(50, y_position, line)
                                y_position -= 15
                    
                    # Clean up temp file
                    temp_image.close()
                    os.unlink(temp_image.name)
                
                c.save()
                QMessageBox.information(self, "Success", "PDF exported successfully!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save PDF: {str(e)}")

    def show_context_menu(self, position):
        menu = QMenu()
        delete_action = menu.addAction("Delete Screenshot")
        
        # Get the item at the position
        item = self.list_widget.itemAt(position)
        
        if item is not None:
            action = menu.exec(self.list_widget.mapToGlobal(position))
            if action == delete_action:
                self.delete_screenshot(self.list_widget.row(item))

    def delete_screenshot(self, index):
        reply = QMessageBox.question(
            self,
            'Delete Screenshot',
            'Are you sure you want to delete this screenshot?',
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # Clear current screenshot index if we're deleting the current screenshot
            if self.current_screenshot_index == index:
                self.current_screenshot_index = None
            
            # Remove from data structures
            self.screenshots.pop(index)
            # Remove from list widget
            self.list_widget.takeItem(index)
            
            # Clear preview if no screenshots left
            if not self.screenshots:
                while self.preview_layout.count():
                    item = self.preview_layout.takeAt(0)
                    if item.widget():
                        item.widget().deleteLater()
                self.notes_edit.clear()
                self.current_screenshot_index = None
            else:
                # Show the previous screenshot if available
                new_index = min(index, len(self.screenshots) - 1)
                self.list_widget.setCurrentRow(new_index)
                self.show_screenshots(self.list_widget.item(new_index))
                self.current_screenshot_index = new_index

def main():
    app = QApplication(sys.argv)
    ex = ScreenshotManager()
    ex.show()
    sys.exit(app.exec())

if __name__ == '__main__':
    main()
